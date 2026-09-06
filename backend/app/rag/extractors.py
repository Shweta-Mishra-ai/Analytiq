"""
rag/extractors.py — turn any supported file into text passages.

Text-ish files are parsed locally. Images and video go through Gemini
(vision + File API) and come back as rich analyst descriptions, so they
become searchable text like everything else.

Every extractor returns a list of passages:
    {"text": str, "source": filename, "locator": "page 3" / "0:45–1:30" / ...}
"""
from __future__ import annotations

import io
import logging
import os
from typing import List


logger = logging.getLogger(__name__)

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
VIDEO_EXTS = {".mp4", ".mov", ".webm", ".avi", ".mkv"}
TEXT_EXTS = {".txt", ".md", ".log"}
TABLE_EXTS = {".csv", ".tsv"}
SHEET_EXTS = {".xlsx", ".xlsm", ".xls"}

IMAGE_PROMPT = """You are a senior data analyst. Analyze this image thoroughly:
1. Describe what the image shows.
2. Transcribe ALL visible text, numbers, labels and table contents exactly.
3. If it contains a chart/graph: name the chart type, axes, series, and read
   out the approximate data values and the trend they show.
4. State 2-3 analytical takeaways.
Be factual — never invent values that are not visible."""

SCAN_PROMPT = """This is a scanned page of a document. Transcribe every
word, number, heading, table cell, signature block, stamp and handwritten
note you can see, in reading order. Preserve table structure using ' | '
between cells. Do not summarise, do not omit small print, and do not
invent anything that is not on the page. If part of the page is illegible,
write [illegible] there rather than guessing."""

VIDEO_PROMPT = """You are a senior data analyst. Analyze this video:
1. Summarize what happens, segment by segment, with timestamps.
2. Transcribe any spoken narration or on-screen text/numbers.
3. If charts, dashboards or data appear: read out the metrics and values shown.
4. State the key analytical takeaways.
Be factual — never invent values that are not shown or said."""


def extract(filename: str, data: bytes) -> List[dict]:
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return _pdf(filename, data)
    if ext == ".docx":
        return _docx(filename, data)
    if ext in TEXT_EXTS:
        return _plain(filename, data)
    if ext in TABLE_EXTS:
        return _table(filename, data)
    if ext in SHEET_EXTS:
        return _spreadsheet(filename, data)
    if ext in IMAGE_EXTS:
        return _image(filename, data)
    if ext in VIDEO_EXTS:
        return _video(filename, data, ext)
    raise ValueError(f"Unsupported file type: {ext}")


# ── documents ────────────────────────────────────────────

# A page carrying only a scanner's image has no text layer. Rendering it
# for a vision model costs a second or two, so it is worth doing for the
# handful of pages in a document that need it and not for a 400-page
# scan somebody uploaded by mistake.
MAX_PAGES_READ_AS_IMAGES = 20

# Below this many characters a page is furniture — a page number, a
# footer — rather than content, and worth reading as an image.
MIN_PAGE_TEXT = 40


def _pdf(name: str, data: bytes) -> List[dict]:
    """Every page accounted for, including the ones that are pictures.

    A page with no text layer used to be skipped in silence. On a wholly
    scanned file that surfaced as "no text could be extracted", which is
    at least honest; but a contract with its signed page photographed and
    reinserted ingested cleanly, reported success, and simply did not
    contain that page. Nothing anywhere said so.

    Such a page is now read by the vision model that already reads
    uploaded images. Where no vision model is configured, the page is
    still recorded — the knowledge base says it holds a page it could
    not read, rather than behaving as though the page did not exist.
    """
    import pypdfium2 as pdfium
    doc = pdfium.PdfDocument(data)
    out, imageish = [], []
    try:
        for i, page in enumerate(doc):
            text = (page.get_textpage().get_text_bounded() or "").strip()
            if len(text) >= MIN_PAGE_TEXT:
                out.append({"text": text, "source": name,
                            "locator": f"page {i + 1}"})
            else:
                imageish.append((i, text))

        for i, leftover in imageish[:MAX_PAGES_READ_AS_IMAGES]:
            read = _page_as_image(doc, i, name)
            if read:
                out.append(read)
            else:
                out.append({
                    "text": ("[Page {} of {} is a scanned image. No text "
                             "model was available to read it, so its "
                             "contents are not searchable.{}]").format(
                                 i + 1, name,
                                 " The page carries the text: " + leftover
                                 if leftover else ""),
                    "source": name, "locator": f"page {i + 1} (unread scan)"})

        if len(imageish) > MAX_PAGES_READ_AS_IMAGES:
            out.append({
                "text": ("[{} of the {} pages in {} are scanned images. The "
                         "first {} were read; the rest are not "
                         "searchable.]").format(
                             len(imageish), len(doc), name,
                             MAX_PAGES_READ_AS_IMAGES),
                "source": name, "locator": "unread scanned pages"})
    finally:
        doc.close()
    # Page order, so a document reads as a document.
    return sorted(out, key=lambda p: _page_no(p["locator"]))


def _page_no(locator: str) -> int:
    import re as _re
    m = _re.search(r"page (\d+)", locator)
    return int(m.group(1)) if m else 10 ** 9


def _page_as_image(doc, index: int, name: str):
    """Render one page and have the vision model read it back.

    Returns None when no vision-capable model is configured, which is a
    normal state — the caller records the gap instead.
    """
    try:
        from app.ai import multimodal
        bitmap = doc[index].render(scale=2)          # ~150 dpi
        buf = io.BytesIO()
        bitmap.to_pil().save(buf, format="PNG")
        text = multimodal.describe_image(
            image=buf.getvalue(), prompt=SCAN_PROMPT,
            task="image_understanding", mime="image/png", timeout_sec=90)
    except Exception:
        # NoCapableModel, a provider error, a render failure: all mean
        # the same thing here, and none of them should lose the page.
        logger.info("page %d of %r could not be read as an image",
                    index + 1, name, exc_info=True)
        return None
    if not (text or "").strip():
        return None
    return {"text": "[Scanned page {} of {}]\n{}".format(
                index + 1, name, text.strip()),
            "source": name, "locator": f"page {index + 1} (scanned)"}


def _docx(name: str, data: bytes) -> List[dict]:
    """Body, tables, and the parts of a document nobody puts in the body.

    Headers and footers were skipped. On internal documentation that is
    where the document reference, the version, the review date and the
    classification live — "Confidential — Internal Use Only", "DOC-4471"
    — and they are exactly what someone asks the knowledge base about.
    """
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))

    furniture = []
    for section in d.sections:
        for area, label in ((section.header, "Header"),
                            (section.footer, "Footer")):
            if area is None:
                continue
            lines = [p.text.strip() for p in area.paragraphs if p.text.strip()]
            for table in getattr(area, "tables", []):
                for row in table.rows:
                    cells = " | ".join(c.text.strip() for c in row.cells)
                    if cells.strip(" |"):
                        lines.append(cells)
            for line in lines:
                entry = "{}: {}".format(label, line)
                # The same header repeats on every section of a document;
                # indexing it once is enough.
                if entry not in furniture:
                    furniture.append(entry)

    out = []
    body = "\n".join(parts).strip()
    if body:
        out.append({"text": body, "source": name, "locator": "document"})
    if furniture:
        out.append({"text": "\n".join(furniture), "source": name,
                    "locator": "header and footer"})
    return out


def _plain(name: str, data: bytes) -> List[dict]:
    text = data.decode("utf-8", errors="replace")
    return [{"text": text, "source": name, "locator": "file"}] if text.strip() else []


# A table is indexed in full, in batches of this many rows. Small enough
# that a retrieved batch is readable and specific; large enough that a
# long log does not become tens of thousands of chunks.
TABLE_ROWS_PER_PASSAGE = 40

# Beyond this the file is a database export, not documentation. It is
# still indexed up to the cap — and the cap is *reported*, never silent.
TABLE_MAX_ROWS = 50_000


def _table(name: str, data: bytes) -> List[dict]:
    """Every row, not a sample of them.

    This used to keep `df.head(50)`. A 5,000-row incident log became 50
    rows and 3,891 characters — 99% of the file discarded, with nothing
    anywhere to say so. Asked about incident 4,000, the knowledge base
    answered from the rows it happened to keep, which is worse than
    admitting it does not know.

    Rows are emitted in batches with the header repeated on each, so a
    retrieved batch stands on its own: a chunk that reads
    `incident_id | system | severity` and then its rows can be quoted
    without the rest of the table.
    """
    import pandas as pd
    sep = "\t" if name.lower().endswith(".tsv") else ","
    df = pd.read_csv(io.BytesIO(data), sep=sep, on_bad_lines="skip")
    return _frame_passages(df, name, "table")


def _frame_passages(df, name: str, locator_kind: str) -> List[dict]:
    """A dataframe as searchable passages: a summary, then all the rows."""
    total = len(df)
    header = " | ".join(str(c) for c in df.columns)

    summary = ["Table {}: {:,} rows, {} columns: {}".format(
        name, total, len(df.columns), ", ".join(map(str, df.columns)))]
    num = df.select_dtypes(include="number")
    if not num.empty:
        summary.append("Numeric summary:\n"
                       + num.describe().round(2).to_string())
    out = [{"text": "\n\n".join(summary), "source": name,
            "locator": "{} summary".format(locator_kind)}]

    body = df.head(TABLE_MAX_ROWS)
    if total > TABLE_MAX_ROWS:
        # Said out loud, in the indexed text itself, so an answer drawn
        # from this table carries the caveat with it.
        out[0]["text"] += (
            "\n\nNOTE: this table has {:,} rows; the first {:,} are "
            "indexed and searchable. Rows beyond that are not included "
            "in answers.".format(total, TABLE_MAX_ROWS))

    for start in range(0, len(body), TABLE_ROWS_PER_PASSAGE):
        batch = body.iloc[start:start + TABLE_ROWS_PER_PASSAGE]
        lines = [" | ".join(str(v) for v in row)
                 for row in batch.itertuples(index=False, name=None)]
        out.append({
            "text": "Rows {:,}–{:,} of {} ({:,} total)\n{}\n{}".format(
                start + 1, start + len(batch), name, total, header,
                "\n".join(lines)),
            "source": name,
            "locator": "{} rows {:,}-{:,}".format(
                locator_kind, start + 1, start + len(batch)),
        })
    return out


def _spreadsheet(name: str, data: bytes) -> List[dict]:
    """Every sheet of a workbook, each indexed like a table.

    A spreadsheet used to be refused outright — "Unsupported file type:
    .xlsx" — which rules out a large share of the documentation anyone
    actually keeps: contract registers, risk logs, budget trackers.
    """
    import pandas as pd
    book = pd.read_excel(io.BytesIO(data), sheet_name=None)
    out = []
    for sheet, frame in book.items():
        if frame.empty:
            continue
        out.extend(_frame_passages(
            frame, "{} [{}]".format(name, sheet), "sheet '{}'".format(sheet)))
    return out


# ── media via Gemini ─────────────────────────────────────

def _image(name: str, data: bytes) -> List[dict]:
    """Describe an image so its content is searchable with the documents.

    Routed rather than pinned to one vendor: any model declaring the
    vision capability can do this, which includes a local multimodal
    model — and for a knowledge base, which holds a client's contracts
    and policies, keeping the option of never leaving the machine is the
    point.
    """
    from app.ai import multimodal
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    img.load()
    try:
        text = multimodal.describe_image(
            image=data, prompt=IMAGE_PROMPT, task="image_understanding",
            mime=Image.MIME.get(img.format or "", "image/png"),
            timeout_sec=60)
    except multimodal.NoCapableModel:
        raise
    except Exception as e:
        raise RuntimeError(f"Image analysis failed: {e}")
    return [{"text": f"[Image analysis of {name}]\n{text}",
             "source": name, "locator": "image"}]


def _video(name: str, data: bytes, ext: str) -> List[dict]:
    """Watch a clip whole — visuals, on-screen text and narration.

    Routed through the video capability rather than named after Gemini,
    even though Gemini is currently the only model that has it. The
    difference shows up the day a second one does: this code needs no
    change, and the System page already says which models can serve it.
    The temp-file and remote-file cleanup both live in ai/multimodal.py.
    """
    from app.ai import multimodal
    try:
        text = multimodal.understand_video(
            data=data, ext=ext, prompt=VIDEO_PROMPT, timeout_sec=120)
    except multimodal.NoCapableModel:
        raise
    except Exception as e:
        raise RuntimeError(f"Video analysis failed: {e}")
    return [{"text": f"[Video analysis of {name}]\n{text}",
             "source": name, "locator": "video"}]
