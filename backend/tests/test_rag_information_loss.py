"""What survives the journey from an uploaded file into the index.

Every number in these tests was measured against the pipeline before it
was changed. A knowledge base built from a company's own documentation
is only worth having if what went in can be got back out; a pipeline
that quietly keeps a fraction of a file is worse than one that refuses
it, because nothing on screen says which fraction.
"""
import io

import pandas as pd
import pytest

from app.rag.extractors import extract
from app.rag.service import chunk_passages


def indexed_text(filename: str, data: bytes) -> str:
    """Everything that would end up searchable for this file."""
    return "\n".join(c["text"] for c in chunk_passages(extract(filename, data)))


# ══════════════════════════════════════════════════════════
#  TABLES
# ══════════════════════════════════════════════════════════

@pytest.fixture()
def incident_log() -> bytes:
    """5,000 rows, each carrying a unique marker."""
    frame = pd.DataFrame({
        "incident_id": [f"INC-{i:05d}" for i in range(5000)],
        "system": [f"system_{i % 40}" for i in range(5000)],
        "severity": [["low", "medium", "high"][i % 3] for i in range(5000)],
        "hours_lost": [round((i % 17) * 0.5, 1) for i in range(5000)],
    })
    return frame.to_csv(index=False).encode()


def test_every_row_of_a_table_is_indexed(incident_log):
    """It kept `df.head(50)`. A 5,000-row incident log became 50 rows and
    3,891 characters — 99% of the file discarded, with nothing anywhere
    to say so."""
    text = indexed_text("incidents.csv", incident_log)
    missing = [f"INC-{i:05d}" for i in range(5000)
               if f"INC-{i:05d}" not in text]
    assert not missing, "{} rows never reached the index, first: {}".format(
        len(missing), missing[:3])


def test_the_last_row_is_there_too(incident_log):
    """The end of a file is where truncation shows first."""
    assert "INC-04999" in indexed_text("incidents.csv", incident_log)


def test_each_batch_carries_its_header(incident_log):
    """A retrieved batch has to stand on its own: rows without column
    names cannot be quoted or checked."""
    passages = extract("incidents.csv", incident_log)
    rows = [p for p in passages if "rows" in p["locator"]]
    assert rows
    for p in rows[:5]:
        assert "incident_id | system | severity" in p["text"]


def test_each_batch_says_where_it_sits_in_the_table(incident_log):
    passages = extract("incidents.csv", incident_log)
    rows = [p for p in passages if "rows" in p["locator"]]
    assert "(5,000 total)" in rows[0]["text"]
    assert rows[0]["locator"] == "table rows 1-40"


def test_the_summary_passage_survives(incident_log):
    """Column names and the numeric summary are what an opening question
    lands on."""
    passages = extract("incidents.csv", incident_log)
    assert passages[0]["locator"] == "table summary"
    assert "5,000 rows" in passages[0]["text"]


def test_a_table_beyond_the_cap_says_so_rather_than_going_quiet(monkeypatch):
    """A cap is defensible. A silent cap is not."""
    from app.rag import extractors
    monkeypatch.setattr(extractors, "TABLE_MAX_ROWS", 100)
    frame = pd.DataFrame({"id": range(250), "v": range(250)})
    passages = extractors.extract("big.csv", frame.to_csv(index=False).encode())
    assert "250" in passages[0]["text"]
    assert "not included" in passages[0]["text"].lower()


# ══════════════════════════════════════════════════════════
#  SPREADSHEETS
# ══════════════════════════════════════════════════════════

@pytest.fixture()
def workbook() -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf) as writer:
        pd.DataFrame({"vendor": [f"vendor_{i}" for i in range(300)],
                      "value": range(300)}).to_excel(
                          writer, sheet_name="Contracts", index=False)
        pd.DataFrame({"risk": [f"risk_{i}" for i in range(50)]}).to_excel(
            writer, sheet_name="Risks", index=False)
    return buf.getvalue()


def test_a_spreadsheet_is_no_longer_refused(workbook):
    """`Unsupported file type: .xlsx` ruled out contract registers, risk
    logs and budget trackers — much of what a company calls its
    documentation."""
    text = indexed_text("contracts.xlsx", workbook)
    assert "vendor_0" in text and "vendor_299" in text


def test_every_sheet_of_a_workbook_is_read(workbook):
    text = indexed_text("contracts.xlsx", workbook)
    assert "risk_49" in text, "the second sheet was not indexed"


def test_a_sheet_is_named_in_the_citation(workbook):
    locators = [p["locator"] for p in extract("contracts.xlsx", workbook)]
    assert any("Contracts" in loc for loc in locators)
    assert any("Risks" in loc for loc in locators)


# ══════════════════════════════════════════════════════════
#  WORD DOCUMENTS
# ══════════════════════════════════════════════════════════

@pytest.fixture()
def policy() -> bytes:
    import docx
    d = docx.Document()
    d.add_heading("Information Security Policy", 0)
    for i in range(1, 21):
        d.add_paragraph(f"Clause {i:03d}. Control CTRL-{i:03d} applies.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Control"
    table.rows[0].cells[1].text = "Owner"
    for i in range(1, 6):
        row = table.add_row()
        row.cells[0].text = f"CTRL-{i:03d}"
        row.cells[1].text = f"owner_{i}"
    section = d.sections[0]
    section.header.paragraphs[0].text = "Confidential — Internal Use Only"
    section.footer.paragraphs[0].text = "Document reference DOC-4471"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_the_body_and_its_tables_survive(policy):
    text = indexed_text("policy.docx", policy)
    assert "Clause 001" in text and "Clause 020" in text
    assert "CTRL-005 | owner_5" in text


def test_the_header_and_footer_survive(policy):
    """This is where the document reference, the version and the
    classification live — and what people ask about."""
    text = indexed_text("policy.docx", policy)
    assert "Confidential — Internal Use Only" in text
    assert "DOC-4471" in text


def test_the_header_is_indexed_once_not_once_per_section(policy):
    text = indexed_text("policy.docx", policy)
    assert text.count("DOC-4471") == 1


# ══════════════════════════════════════════════════════════
#  PDFS, INCLUDING THE PAGES THAT ARE PICTURES
# ══════════════════════════════════════════════════════════

def _text_pdf() -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate
    buf = io.BytesIO()
    st = getSampleStyleSheet()
    SimpleDocTemplate(buf, pagesize=A4).build(
        [Paragraph(f"Section {i:03d}. Payment term PT-{i:03d} applies.",
                   st["Normal"]) for i in range(1, 31)])
    return buf.getvalue()


def _scan_pdf() -> bytes:
    from PIL import Image, ImageDraw
    img = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(img)
    for i in range(20):
        draw.text((80, 80 + i * 40), f"SCANNED LINE {i:03d}", fill="black")
    buf = io.BytesIO()
    img.save(buf, "PDF", resolution=150)
    return buf.getvalue()


@pytest.fixture()
def mixed_pdf() -> bytes:
    """A contract with its signed page photographed and reinserted."""
    import pypdf
    writer = pypdf.PdfWriter()
    writer.add_page(pypdf.PdfReader(io.BytesIO(_text_pdf())).pages[0])
    writer.add_page(pypdf.PdfReader(io.BytesIO(_scan_pdf())).pages[0])
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_a_scanned_page_is_never_dropped_in_silence(mixed_pdf):
    """It ingested cleanly, reported success, and simply did not contain
    the page. Nothing anywhere said so."""
    passages = extract("mixed.pdf", mixed_pdf)
    assert len(passages) == 2, [p["locator"] for p in passages]
    assert any("2" in p["locator"] for p in passages)


def test_the_unread_page_says_why_it_is_unreadable(mixed_pdf):
    passages = extract("mixed.pdf", mixed_pdf)
    scan = next(p for p in passages if "2" in p["locator"])
    assert "scanned image" in scan["text"]
    assert "not searchable" in scan["text"]


def test_a_scanned_page_is_read_when_a_vision_model_exists(
        mixed_pdf, monkeypatch):
    from app.ai import multimodal
    monkeypatch.setattr(
        multimodal, "describe_image",
        lambda **kw: "SCANNED LINE 000\nSCANNED LINE 001")
    text = indexed_text("mixed.pdf", mixed_pdf)
    assert "SCANNED LINE 000" in text


def test_pages_come_back_in_page_order(mixed_pdf, monkeypatch):
    """A document that reads out of order is a different document."""
    from app.ai import multimodal
    monkeypatch.setattr(multimodal, "describe_image",
                        lambda **kw: "the scanned page")
    locators = [p["locator"] for p in extract("mixed.pdf", mixed_pdf)]
    assert locators == ["page 1", "page 2 (scanned)"]


def test_the_text_pages_are_unaffected(mixed_pdf):
    text = indexed_text("mixed.pdf", mixed_pdf)
    assert "Section 001" in text and "PT-030" in text


# ══════════════════════════════════════════════════════════
#  WHAT THE ANSWER IS ALLOWED TO CLAIM
# ══════════════════════════════════════════════════════════

def _row_hit(source, first, last, total):
    return {
        "source": source,
        "locator": "table rows {}-{}".format(first, last),
        "text": "Rows {}–{} of {} ({:,} total)\nid | v\n{} | 1".format(
            first, last, source, total, first),
    }


def test_a_partly_shown_table_is_declared_as_partial():
    """Search returns the passages closest to the question. For a table
    that is six batches of forty rows out of five thousand, and "how many
    incidents were high severity" then has an exact answer that cannot be
    computed from what is in front of the model. A precise number from a
    fraction of the rows is the most damaging answer available."""
    from app.rag.service import _context_block

    block = _context_block([_row_hit("incidents.csv", 1, 40, 5000),
                            _row_hit("incidents.csv", 201, 240, 5000)])
    assert "COVERAGE OF TABLES" in block
    assert "80 of 5,000 rows" in block


def test_a_fully_shown_table_gets_no_such_warning():
    """The note must mean something when it appears."""
    from app.rag.service import _context_block
    assert "COVERAGE" not in _context_block(
        [_row_hit("small.csv", 1, 40, 40)])


def test_prose_passages_never_trigger_the_warning():
    from app.rag.service import _context_block
    block = _context_block([{"source": "policy.docx", "locator": "document",
                             "text": "Clause 001. Control CTRL-001 applies."}])
    assert "COVERAGE" not in block


def test_two_tables_are_reported_separately():
    from app.rag.service import _context_block
    block = _context_block([_row_hit("a.csv", 1, 40, 900),
                            _row_hit("b.csv", 1, 40, 4000)])
    assert "a.csv: 40 of 900 rows" in block
    assert "b.csv: 40 of 4,000 rows" in block


def test_the_model_is_told_not_to_total_from_a_partial_table():
    """The context can only carry the caveat; the instruction has to
    forbid the arithmetic."""
    from app.rag.service import QA_SYSTEM
    lowered = QA_SYSTEM.lower()
    assert "coverage" in lowered
    assert "do not state a count" in lowered
